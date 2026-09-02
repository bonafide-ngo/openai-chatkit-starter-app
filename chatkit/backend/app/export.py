from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import re
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from xml.sax.saxutils import escape


EXPORT_TEXT = {
    "en": {"chat_export": "Chat export", "title": "Title", "exported": "Exported", "user": "User", "assistant": "Assistant", "message": "Message"},
    "de": {"chat_export": "Chat-Export", "title": "Titel", "exported": "Exportiert", "user": "Benutzer", "assistant": "Assistent", "message": "Nachricht"},
    "es": {"chat_export": "Exportación del chat", "title": "Título", "exported": "Exportado", "user": "Usuario", "assistant": "Asistente", "message": "Mensaje"},
    "fr": {"chat_export": "Exportation de la conversation", "title": "Titre", "exported": "Exporté", "user": "Utilisateur", "assistant": "Assistant", "message": "Message"},
    "it": {"chat_export": "Esportazione della chat", "title": "Titolo", "exported": "Esportato", "user": "Utente", "assistant": "Assistente", "message": "Messaggio"},
    "ja": {"chat_export": "チャットのエクスポート", "title": "タイトル", "exported": "エクスポート日時", "user": "ユーザー", "assistant": "アシスタント", "message": "メッセージ"},
    "ko": {"chat_export": "채팅 내보내기", "title": "제목", "exported": "내보낸 시간", "user": "사용자", "assistant": "어시스턴트", "message": "메시지"},
    "nl": {"chat_export": "Chatexport", "title": "Titel", "exported": "Geëxporteerd", "user": "Gebruiker", "assistant": "Assistent", "message": "Bericht"},
    "pl": {"chat_export": "Eksport rozmowy", "title": "Tytuł", "exported": "Wyeksportowano", "user": "Użytkownik", "assistant": "Asystent", "message": "Wiadomość"},
    "pt": {"chat_export": "Exportação da conversa", "title": "Título", "exported": "Exportado", "user": "Usuário", "assistant": "Assistente", "message": "Mensagem"},
    "ru": {"chat_export": "Экспорт разговора", "title": "Название", "exported": "Экспортировано", "user": "Пользователь", "assistant": "Ассистент", "message": "Сообщение"},
    "zh": {"chat_export": "聊天导出", "title": "标题", "exported": "导出时间", "user": "用户", "assistant": "助手", "message": "消息"},
}


def export_text(locale: str) -> dict[str, str]:
    return EXPORT_TEXT.get(locale, EXPORT_TEXT["en"])


def _text(value: Any) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return " ".join(part for item in value if (part := _text(item)))
    if isinstance(value, dict):
        if value.get("type") in {"input_text", "output_text", "text"}:
            return _text(value.get("text", ""))
        return " ".join(part for key, value in value.items() if key not in {"annotations", "file_data"} if (part := _text(value)))
    return ""


def item_text(item: Any) -> str:
    data = item.model_dump(mode="json")
    return _text(data.get("content", data.get("text", ""))).strip()


def item_role(item: Any) -> str:
    item_type = getattr(item, "type", "")
    if "user" in item_type or "input" in item_type:
        return "User"
    if "assistant" in item_type or "response" in item_type:
        return "Assistant"
    return item_type.replace("_", " ").title() or "Message"


CODE_BLOCK_PATTERN = re.compile(r"```[^\n]*\n?(.*?)```", re.DOTALL)


def content_segments(text: str) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    cursor = 0
    for match in CODE_BLOCK_PATTERN.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor : match.start()], False))
        segments.append((match.group(1).rstrip("\n"), True))
        cursor = match.end()
    if cursor < len(text):
        segments.append((text[cursor:], False))
    return [(content, is_code) for content, is_code in segments if content.strip()]


def conversation_rows(thread: Any, items: list[Any], locale: str = "en") -> list[tuple[str, str]]:
    labels = export_text(locale)
    rows = []
    for item in items:
        text = item_text(item)
        if text:
            role = item_role(item)
            role = labels.get(role.lower(), role)
            rows.append((role, text))
    return rows


def build_docx(app_title: str, title: str, rows: list[tuple[str, str]], exported_at: datetime, locale: str = "en") -> BytesIO:
    labels = export_text(locale)
    document = Document()
    app_heading = document.add_heading(app_title, level=0)
    for run in app_heading.runs:
        run.font.underline = False
    metadata = document.add_paragraph()
    title_run = metadata.add_run(f"{labels['title']}: {title}")
    title_run.bold = True
    title_run.font.underline = False
    exported_run = metadata.add_run(
        f"\n{labels['exported']}: {exported_at.strftime('%Y-%m-%d %H:%M UTC')}"
    )
    exported_run.font.underline = False
    code_style = document.styles.add_style("ExportCode", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    for role, text in rows:
        heading = document.add_heading(role, level=2)
        for run in heading.runs:
            run.font.underline = False
            run.font.color.rgb = RGBColor.from_string(
                "2563EB" if role == labels["user"] else "D97706" if role == labels["assistant"] else "64748B"
            )
        for content, is_code in content_segments(text):
            paragraph = document.add_paragraph(content, style="ExportCode" if is_code else None)
            paragraph.paragraph_format.space_after = 8
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def build_pdf(app_title: str, title: str, rows: list[tuple[str, str]], exported_at: datetime, locale: str = "en") -> BytesIO:
    labels = export_text(locale)
    output = BytesIO()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="AppTitle",
        parent=styles["Title"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="ChatTitle",
        parent=styles["Title"],
        fontSize=22,
        leading=27,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="Metadata",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=16,
    ))
    styles.add(ParagraphStyle(
        name="UserRole",
        parent=styles["Heading2"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#2563eb"),
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="AssistantRole",
        parent=styles["Heading2"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#d97706"),
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="OtherRole",
        parent=styles["Heading2"],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#64748b"),
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CodeBlock",
        parent=styles["BodyText"],
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        backColor=colors.HexColor("#f1f5f9"),
        borderColor=colors.HexColor("#cbd5e1"),
        borderWidth=0.5,
        borderPadding=6,
        spaceAfter=8,
    ))
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph(escape(app_title), styles["AppTitle"]),
        Paragraph(escape(title), styles["ChatTitle"]),
        Paragraph(
            f"{escape(labels['exported'])} {escape(exported_at.strftime('%Y-%m-%d %H:%M UTC'))}",
            styles["Metadata"],
        ),
    ]
    for role, text in rows:
        role_style = "UserRole" if role == labels["user"] else "AssistantRole" if role == labels["assistant"] else "OtherRole"
        story.extend(
            [
                Paragraph(escape(role), styles[role_style]),
            ]
        )
        for content, is_code in content_segments(text):
            formatted_content = escape(content).replace(" ", "&nbsp;").replace("\n", "<br/>")
            story.extend([
                Paragraph(formatted_content, styles["CodeBlock" if is_code else "BodyText"]),
                Spacer(1, 2 * mm),
            ])
    document.build(story)
    output.seek(0)
    return output


def build_markdown(app_title: str, title: str, rows: list[tuple[str, str]], exported_at: datetime, locale: str = "en") -> BytesIO:
    labels = export_text(locale)
    markdown = [
        f"# {app_title}",
        f"## {title}",
        f"**{labels['exported']}:** {exported_at.strftime('%Y-%m-%d %H:%M UTC')}",
        "",
    ]
    for role, text in rows:
        markdown.extend([f"### {role}", "", text, ""])
    output = BytesIO("\n".join(markdown).encode("utf-8"))
    output.seek(0)
    return output
